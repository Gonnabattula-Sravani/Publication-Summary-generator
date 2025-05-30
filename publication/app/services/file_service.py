from docx import Document
import xlsxwriter

def generate_docx(publications):
    doc = Document()
    doc.add_heading('Publications Overview', 0)

    for pub in publications:
        faculty_name = pub.get('faculty_name', 'Unknown Faculty')
        title = pub.get('title', 'Unknown Title')
        year = pub.get('year', 'Unknown Year')
        pub_type = pub.get('type', 'Unknown Type')

        doc.add_paragraph(f"Faculty Name: {faculty_name}")
        doc.add_paragraph(f"Title: {title}")
        doc.add_paragraph(f"Year: {year}")
        doc.add_paragraph(f"Type: {pub_type}")
        doc.add_paragraph(f"Citation Count: {pub['citation_count']}")
        doc.add_paragraph("\n")  # Add spacing between entries

    filename = 'all_publications.docx'  # A single file for all faculty
    doc.save(filename)
    return filename





def generate_xlsx(publications):
    filepath = 'all_publications.xlsx'  # A single file for all faculty
    workbook = xlsxwriter.Workbook(filepath)
    worksheet = workbook.add_worksheet()

    # Write headers including 'Faculty Name', 'Title', 'Year', 'Type', and 'Citation Count'
    worksheet.write(0, 0, 'Faculty Name')
    worksheet.write(0, 1, 'Title')
    worksheet.write(0, 2, 'Year')
    worksheet.write(0, 3, 'Type')  # Adding the type field
    worksheet.write(0, 4, 'Citation Count')  # Adding citation count field

    # Adjust column widths for readability
    worksheet.set_column(0, 0, 20)  # Adjusting column width for Faculty Name
    worksheet.set_column(1, 1, 50)  # Set column width for Title
    worksheet.set_column(2, 2, 15)  # Set column width for Year
    worksheet.set_column(3, 3, 20)  # Set column width for Type
    worksheet.set_column(4, 4, 15)  # Set column width for Citation Count

    # Write publication data starting from row 1
    for i, pub in enumerate(publications, start=1):
        # Use get() to safely fetch 'faculty_name', 'title', 'year', 'type', and 'citation_count' with default values if missing
        faculty_name = pub.get('faculty_name', 'Unknown Faculty')
        title = pub.get('title', 'Unknown Title')
        year = pub.get('year', 'Unknown Year')
        pub_type = pub.get('type', 'Unknown Type')
        citation_count = pub.get('citation_count', 0)  # Default to 0 if citation_count is missing

        # Write each row of data
        worksheet.write(i, 0, faculty_name)
        worksheet.write(i, 1, title)
        worksheet.write(i, 2, year)
        worksheet.write(i, 3, pub_type)
        worksheet.write(i, 4, citation_count)  # Write citation count

    workbook.close()
    return filepath

